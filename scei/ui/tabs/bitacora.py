
from datetime import datetime, timedelta
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLineEdit, QTableWidget, QTableWidgetItem,
    QHeaderView, QPushButton, QLabel, QMessageBox, QFileDialog, QDialog
)
from PyQt6.QtCore import Qt, QSettings, QTimer

from ...utils import load_icon, export_table_to_excel, NumericTableWidgetItem
from ...logger import LOGS, add_log, save_logs # Kept for compat, but we will use DB mainly now
from ...config import BITACORA_CLEAN_INTERVAL_DAYS
from ..dialogs import RecordDetailDialog, AdminAuthDialog
from ...data.repositories import list_bitacora_entries, add_bitacora_log, get_user
from ... import session

class BitacoraTab(QWidget):
    def __init__(self):
        super().__init__()
        self.settings = QSettings("SCEI", "App")
        self._bitacora_last_cleanup: datetime | None = None

        # Layout Principal (sin márgenes excesivos para aprovechar espacio)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(32, 24, 32, 32)
        layout.setSpacing(24)

        # --- Columna Izquierda: Acciones ---
        left_card = QWidget()
        left_card.setProperty("class", "card")
        left_card.setFixedWidth(200) # Un poco más ancho para que los botones respiren
        left_l = QVBoxLayout(left_card)
        left_l.setContentsMargins(16, 20, 16, 20)
        left_l.setSpacing(12)

        lbl_actions = QLabel("Acciones")
        lbl_actions.setStyleSheet("font-weight: 700; color: #94A3B8; text-transform: uppercase; font-size: 12px; margin-bottom: 8px;")
        left_l.addWidget(lbl_actions)

        btn_pdf = QPushButton("Exportar PDF")
        btn_pdf.setIcon(load_icon("pdf.svg")) # Asumiendo existencia, si no carga empty
        btn_word = QPushButton("Exportar Word")
        btn_word.setIcon(load_icon("word.svg"))
        btn_excel = QPushButton("Exportar Excel")
        btn_excel.setIcon(load_icon("excel.svg"))
        
        self.btn_clear = QPushButton("Limpiar Historial")
        self.btn_clear.setEnabled(False)
        self.btn_clear.setProperty("class", "panel-accent")
        self.btn_clear.setMinimumHeight(42)
        self.btn_clear.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_clear.setStyleSheet("QPushButton { color: #EF4444; border: 1px solid rgba(239, 68, 68, 0.3); } QPushButton:hover { background: rgba(239, 68, 68, 0.1); border: 1px solid rgba(239, 68, 68, 0.8); }")

        for b in (btn_pdf, btn_word, btn_excel):
            b.setProperty("class", "panel-accent")
            b.setMinimumHeight(42) # Botones más altos
            b.setCursor(Qt.CursorShape.PointingHandCursor)
            left_l.addWidget(b)
        
        left_l.addStretch(1)
        left_l.addWidget(self.btn_clear)
        
        layout.addWidget(left_card)

        # --- Columna Derecha: Contenido ---
        right_layout = QVBoxLayout()
        right_layout.setSpacing(16)

        # Header y Buscador (En una fila transparente)
        header_row = QHBoxLayout()
        lbl_title = QLabel("Bitácora de Actividades")
        lbl_title.setStyleSheet("font-size: 24px; font-weight: 700; color: #F8FAFC;")
        
        self.search = QLineEdit()
        self.search.setPlaceholderText("Buscar en registros...")
        self.search.setFixedWidth(300)
        self.search.setStyleSheet("""
            QLineEdit {
                border-radius: 18px;
                padding: 8px 16px;
                background: rgba(30, 41, 59, 0.6);
                border: 1px solid rgba(148, 163, 184, 0.2);
                color: #F1F5F9;
            }
            QLineEdit:focus {
                border: 1px solid #38BDF8;
                background: rgba(30, 41, 59, 0.9);
            }
        """)
        
        header_row.addWidget(lbl_title)
        header_row.addStretch(1)
        header_row.addWidget(self.search)
        right_layout.addLayout(header_row)

        # Tarjeta de Tabla
        table_card = QWidget()
        table_card.setProperty("class", "card")
        table_l = QVBoxLayout(table_card)
        table_l.setContentsMargins(0, 0, 0, 0) # La tabla llena la tarjeta, padding en celdas
        
        self.table = QTableWidget(0, 5)
        self.table = QTableWidget(0, 5)
        self.table.setHorizontalHeaderLabels(["N°", "Fecha", "Usuario", "Acción", "Descripción"])
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setShowGrid(False) # Estilo premium
        self.table.verticalHeader().setVisible(False)
        self.table.setFrameShape(QTableWidget.Shape.NoFrame)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        # Ajustar anchos específicos
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents) # N°
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents) # Fecha
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents) # Usuario
        
        table_l.addWidget(self.table)
        right_layout.addWidget(table_card, 1)

        layout.addLayout(right_layout, 1)

        # Conexiones
        self.search.textChanged.connect(self.refresh)
        btn_pdf.clicked.connect(self.generar_pdf)
        btn_word.clicked.connect(self.generar_word)
        btn_excel.clicked.connect(self.generar_excel)
        self.btn_clear.clicked.connect(self.on_clear)
        self.table.itemDoubleClicked.connect(self.show_detail)
        # Auto-renumerar al ordenar
        self.table.horizontalHeader().sectionClicked.connect(self.schedule_renumber)
        self.table.horizontalHeader().sortIndicatorChanged.connect(self.schedule_renumber)

        self.refresh()

    def schedule_renumber(self):
        QTimer.singleShot(50, self.update_row_numbers)

    def update_row_numbers(self):
        if self.table.model(): self.table.model().blockSignals(True)
        for r in range(self.table.rowCount()):
            item = self.table.item(r, 0)
            if not item:
                item = NumericTableWidgetItem(str(r + 1))
                self.table.setItem(r, 0, item)
            else:
                item.setText(str(r + 1))
            item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        if self.table.model(): self.table.model().blockSignals(False)

    def refresh(self):
        self._update_clear_button()
        self.table.setRowCount(0)
        
        # Load from DB instead of memory LOGS
        logs_db = list_bitacora_entries(limit=100) # Limit logic or pagination ideally
        
        term = self.search.text().strip().lower()
        for log in logs_db:
            username = log.usuario.username if log.usuario else "Sistema"
            if term:
                text = f"{log.fecha} {log.accion} {log.descripcion} {username}".lower()
                if term not in text:
                    continue
            
            r = self.table.rowCount()
            self.table.insertRow(r)
            item_n = NumericTableWidgetItem(str(r + 1))
            item_n.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(r, 0, item_n)
            
            # Format date friendly
            d_str = log.fecha.strftime("%Y-%m-%d %H:%M") if log.fecha else ""
            
            self.table.setItem(r, 1, QTableWidgetItem(d_str))
            self.table.setItem(r, 2, QTableWidgetItem(username))
            self.table.setItem(r, 3, QTableWidgetItem(log.accion))
            self.table.setItem(r, 4, QTableWidgetItem(log.descripcion))
        
        self.update_row_numbers()

    def show_detail(self):
        r = self.table.currentRow()
        if r < 0:
            return
        get = lambda c: (self.table.item(r, c).text() if self.table.item(r, c) else "")
        details = [
            ("Fecha", get(1)),
            ("Usuario", get(2)),
            ("Acción", get(3)),
            ("Descripción", get(4)),
        ]
        RecordDetailDialog("Detalle de Bitácora", details).exec()

    def _load_last_cleanup(self) -> datetime | None:
        value = self.settings.value("bitacora_last_cleanup", "")
        if value:
            try:
                return datetime.fromisoformat(str(value))
            except:
                return None
        return None

    def _is_cleanup_due(self) -> bool:
        last = self._load_last_cleanup()
        if last is None:
            return True
        return datetime.now() - last >= timedelta(days=BITACORA_CLEAN_INTERVAL_DAYS)

    def _days_until_cleanup(self) -> int:
        last = self._load_last_cleanup()
        if last is None:
            return 0
        delta = datetime.now() - last
        interval = timedelta(days=BITACORA_CLEAN_INTERVAL_DAYS)
        if delta >= interval:
            return 0
        return int(((interval - delta).total_seconds() + 86399) // 86400)

    def _update_clear_button(self):
        due = self._is_cleanup_due()
        self.btn_clear.setEnabled(due)
        if due:
            self.btn_clear.setToolTip("Limpiar bitácora mensual.")
        else:
            self.btn_clear.setToolTip(f"Disponible en {self._days_until_cleanup()} días.")

    def on_clear(self):
        # Restriction: Only Admin or Authed
        if session.CURRENT_USER != "DI-ADMIN":
             auth = AdminAuthDialog(self, "Se requieren permisos de administrador para limpiar el historial.")
             if auth.exec() != QDialog.DialogCode.Accepted:
                 return

        if not self._is_cleanup_due():
            return
        if QMessageBox.question(self, "Confirmar", "¿Eliminar todos los registros?", 
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.Yes:
            global LOGS
            LOGS.clear()
            save_logs()
            self.settings.setValue("bitacora_last_cleanup", datetime.now().isoformat())
            self.settings.sync()
            add_log("Limpiar Bitácora", "Reinicio mensual")
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Limpiar Historial", "Se realizó limpieza de registros", "Bitacora")
            except: pass
            
            self.refresh()

    # --- Generación de Reportes (Simplificada) ---
    def generar_pdf(self):
        if session.CURRENT_USER != "DI-ADMIN":
             auth = AdminAuthDialog(self, "Se requieren permisos de administrador para exportar datos.")
             if auth.exec() != QDialog.DialogCode.Accepted:
                 return

        from PyQt6.QtPrintSupport import QPrinter
        from PyQt6.QtGui import QTextDocument
        html = "<html><body><h1>Bitácora</h1><table border='1' width='100%'>"
        html += "<tr><th>Fecha</th><th>Usuario</th><th>Acción</th><th>Descripción</th></tr>"
        # Export logic should probably fetch fresh data too, but for simplicity using table dump
        for r in range(self.table.rowCount()):
             get = lambda c: (self.table.item(r, c).text() if self.table.item(r, c) else "")
             html += f"<tr><td>{get(1)}</td><td>{get(2)}</td><td>{get(3)}</td><td>{get(4)}</td></tr>"
        html += "</table></body></html>"
        doc = QTextDocument(); doc.setHtml(html)
        printer = QPrinter(); printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        fn, _ = QFileDialog.getSaveFileName(self, "PDF", "bitacora.pdf", "PDF (*.pdf)")
        if fn:
            printer.setOutputFileName(fn)
            doc.print(printer)
            add_log("Generar PDF Bitácora", f"{fn}")
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar PDF Bitácora", f"Archivo: {fn}", "Bitacora")
            except: pass
            
            self.refresh()
            QMessageBox.information(self, "PDF Generado", f"Archivo guardado correctamente en:\n{fn}")

    def generar_word(self):
        if session.CURRENT_USER != "DI-ADMIN":
             auth = AdminAuthDialog(self, "Se requieren permisos de administrador para exportar datos.")
             if auth.exec() != QDialog.DialogCode.Accepted:
                 return

        try:
            from docx import Document
        except:
            QMessageBox.warning(self, "Error", "Falta python-docx")
            return
        doc = Document()
        doc.add_heading('Bitácora', 0)
        table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
        headers = ['Fecha','Usuario','Acción','Descripción']
        for i,h in enumerate(headers): table.rows[0].cells[i].text = h
        for r in range(self.table.rowCount()):
            row = table.add_row().cells
            get = lambda c: (self.table.item(r, c).text() if self.table.item(r, c) else "")
            row[0].text = get(1)
            row[1].text = get(2)
            row[2].text = get(3)
            row[3].text = get(4)
        fn, _ = QFileDialog.getSaveFileName(self, "Word", "bitacora.docx", "Word (*.docx)")
        if fn:
            doc.save(fn)
            add_log("Generar Word Bitácora", f"{fn}")
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar Word Bitácora", f"Archivo: {fn}", "Bitacora")
            except: pass
            
            self.refresh()
            QMessageBox.information(self, "Word Generado", f"Archivo guardado correctamente en:\n{fn}")

    def generar_excel(self):
        if session.CURRENT_USER != "DI-ADMIN":
             auth = AdminAuthDialog(self, "Se requieren permisos de administrador para exportar datos.")
             if auth.exec() != QDialog.DialogCode.Accepted:
                 return

        # Use existing utility just by passing the main table
        fn, _ = QFileDialog.getSaveFileName(self, "Excel", "bitacora.xlsx", "Excel (*.xlsx)")
        if fn:
            export_table_to_excel(self.table, fn)
            add_log("Generar Excel Bitácora", f"{fn}")
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar Excel Bitácora", f"Archivo: {fn}", "Bitacora")
            except: pass
            
            self.refresh()
            QMessageBox.information(self, "Excel Generado", f"Archivo guardado correctamente en:\n{fn}")

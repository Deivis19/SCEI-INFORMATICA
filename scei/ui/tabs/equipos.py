
from datetime import date
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton,
    QTableWidget, QTableWidgetItem, QHeaderView, QComboBox, QDateEdit,
    QMenu, QFileDialog, QMessageBox, QDialog
)
from PyQt6.QtCore import Qt, QDate, QTimer

from ...utils import load_icon, export_table_to_excel, NumericTableWidgetItem
from ...logger import add_log
from ... import session
from ..helpers import direccion_nombre
from ..dialogs import GenerateDialog, EquipoDialog, RecordDetailDialog, AdminAuthDialog
from ...data.repositories import (
    list_direcciones, list_equipos, list_equipos_by_direccion, 
    add_equipo, update_equipo, delete_equipo, get_equipo,
    add_bitacora_log, get_user
)
from sqlalchemy.exc import IntegrityError

class EquiposTab(QWidget):
    def __init__(self, departamento_id: int | None = None, direccion_id: int | None = None):
        super().__init__()
        # Departamento eliminado del modelo; ignoramos cualquier valor entrante
        self.departamento_filter = None
        self.direccion_filter = direccion_id
        self.table = QTableWidget(0, 7)
        self.table.setHorizontalHeaderLabels([
            "N°", "Código","Descripción","Marca","Modelo","Serie","Estado"
        ])
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setAlternatingRowColors(True)
        self.table.setShowGrid(False)  # Desactivar grid nativo para estilo CSS premium
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents) # Columna N° ajustada
        self.table.setSortingEnabled(True)
        self.search = QLineEdit(); self.search.setPlaceholderText("Buscar por código/desc/modelo/serie…")
        self.codigo = QLineEdit()
        self.desc = QLineEdit()
        self.marca = QLineEdit()
        self.modelo = QLineEdit()
        self.serie = QLineEdit()
        self.estado = QComboBox(); self.estado.addItems(["optimo","defectuoso","inoperativo"])
        self.dir = QComboBox()
        self.fecha = QDateEdit(); self.fecha.setDate(QDate.currentDate())

        self.btn_add = QPushButton("Nuevo")
        self.btn_upd = QPushButton("Editar")
        self.btn_del = QPushButton("Eliminar")
        self.btn_add.clicked.connect(self.on_add_modal)
        self.btn_upd.clicked.connect(self.on_edit_modal)
        self.btn_del.clicked.connect(self.on_delete)

        # Layout Principal
        layout = QHBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 24)
        layout.setSpacing(20)

        # --- Panel Izquierdo: Acciones ---
        actions_card = QWidget()
        actions_card.setProperty("class", "card")
        actions_card.setFixedWidth(200)
        
        actions_l = QVBoxLayout(actions_card)
        actions_l.setContentsMargins(16, 20, 16, 20)
        actions_l.setSpacing(12)
        
        lbl_actions = QLabel("Acciones")
        lbl_actions.setStyleSheet("font-weight: 700; color: #94A3B8; text-transform: uppercase; font-size: 12px; margin-bottom: 8px;")
        actions_l.addWidget(lbl_actions)

        # Botones
        for btn, label, icon in [
            (self.btn_add, "Nuevo Equipo", "plus.svg"),
            (self.btn_upd, "Editar Selecc.", "edit.svg"),
            (self.btn_del, "Eliminar Selecc.", "trash.svg")
        ]:
            btn.setText(label)
            btn.setIcon(load_icon(icon))
            btn.setProperty("class", "panel-accent")
            btn.setMinimumHeight(42)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            # Estilo específico eliminar
            if btn == self.btn_del:
                 btn.setStyleSheet("QPushButton { color: #EF4444; border: 1px solid rgba(239, 68, 68, 0.3); } QPushButton:hover { background: rgba(239, 68, 68, 0.1); border: 1px solid rgba(239, 68, 68, 0.8); }")
            actions_l.addWidget(btn)

        gen_btn = QPushButton("Generar Reporte")
        gen_btn.setProperty("class", "panel-accent")
        gen_btn.setMinimumHeight(42)
        gen_btn.setIcon(load_icon("records.svg"))
        gen_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        gen_btn.clicked.connect(self.on_generate)
        
        actions_l.addStretch(1)
        actions_l.addWidget(gen_btn)
        
        layout.addWidget(actions_card)

        # --- Panel Derecho: Contenido ---
        content_layout = QVBoxLayout()
        content_layout.setSpacing(16)
        
        # Header Row (Titulo + Buscador)
        header_row = QHBoxLayout()
        
        self.header = QLabel("Gestión de Equipos")
        self.header.setStyleSheet("font-size: 24px; font-weight: 700; color: #F8FAFC;")
        
        self.search.setPlaceholderText("Buscar equipos...")
        self.search.setFixedWidth(300)
        self.search.setStyleSheet("""
            QLineEdit {
                border-radius: 18px;
                padding: 8px 16px;
                background: rgba(30, 41, 59, 0.6);
                border: 1px solid rgba(148, 163, 184, 0.2);
                color: #F1F5F9;
            }
            QLineEdit:focus {
                border: 1px solid #38BDF8;
                background: rgba(30, 41, 59, 0.9);
            }
        """)

        header_row.addWidget(self.header)
        header_row.addStretch(1)
        header_row.addWidget(self.search)
        
        content_layout.addLayout(header_row)

        # Tabla Card
        table_card = QWidget()
        table_card.setProperty("class", "card")
        table_l = QVBoxLayout(table_card)
        table_l.setContentsMargins(0, 0, 0, 0)
        
        # Tabla Config
        self.table.setShowGrid(False)
        self.table.verticalHeader().setVisible(False)
        self.table.setFrameShape(QTableWidget.Shape.NoFrame)
        
        table_l.addWidget(self.table)
        content_layout.addWidget(table_card, 1)

        layout.addLayout(content_layout, 1)
        
        # Init logic
        self.refresh_refs()
        self.refresh()
        self._update_header()
        self.search.textChanged.connect(self.refresh)
        self.table.itemDoubleClicked.connect(self.on_edit_modal)
        # Context menu para generar desde selección
        self.table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.on_table_context_menu)
        
        # Auto-renumerar al ordenar (usando Timer para asegurar que ocurra post-sort)
        # layoutChanged a veces no se dispara en sorting simple de QTableWidget o es tricky
        self.table.horizontalHeader().sectionClicked.connect(self.schedule_renumber)
        self.table.horizontalHeader().sortIndicatorChanged.connect(self.schedule_renumber)

    def schedule_renumber(self):
        QTimer.singleShot(50, self.update_row_numbers)

    def update_row_numbers(self):
        # Bloquear señales del MODELO para evitar trigger de layoutChanged/dataChanged recursivo
        if self.table.model():
            self.table.model().blockSignals(True)
        
        for r in range(self.table.rowCount()):
            item = self.table.item(r, 0)
            if not item:
                item = NumericTableWidgetItem(str(r + 1))
                self.table.setItem(r, 0, item)
            else:
                # Asegurar que sea NumericTableWidgetItem si ya existía (o actualizar texto)
                item.setText(str(r + 1))
                
            item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            
        if self.table.model():
            self.table.model().blockSignals(False)

    def refresh_refs(self):
        self.dir.clear()
        dirs = list_direcciones()
        self.dir.addItem("", None)
        for d in dirs:
            self.dir.addItem(d.nombre, d.id)
        # Lock selection if filtered by Dirección
        if self.direccion_filter:
            idx = self.dir.findData(self.direccion_filter)
            if idx >= 0:
                self.dir.setCurrentIndex(idx)
            self.dir.setEnabled(False)
        self._update_header()

    def refresh(self):
        if self.direccion_filter:
            data = list_equipos_by_direccion(self.direccion_filter)
        else:
            data = list_equipos()
        # Evitar glitches con ordenamiento al insertar filas
        was_sorting = self.table.isSortingEnabled()
        if was_sorting:
            self.table.setSortingEnabled(False)
        self.table.setRowCount(0)
        term = self.search.text().strip().lower()
        for e in data:
            dir_name = direccion_nombre(getattr(e, 'direccion_id', None)) or ""
            text_blob = " ".join([
                e.codigo_interno or "",
                e.descripcion or "",
                e.marca or "",
                e.modelo or "",
                e.nro_serie or "",
                e.estado or "",
                dir_name,
            ]).lower()
            if term and term not in text_blob:
                continue
            r = self.table.rowCount()
            self.table.insertRow(r)
            # Columna N°
            item_n = NumericTableWidgetItem(str(r + 1))
            item_n.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(r, 0, item_n)
            self.table.setItem(r, 1, QTableWidgetItem(e.codigo_interno or ""))
            self.table.setItem(r, 2, QTableWidgetItem(e.descripcion or ""))
            self.table.setItem(r, 3, QTableWidgetItem(e.marca or ""))
            self.table.setItem(r, 4, QTableWidgetItem(e.modelo or ""))
            self.table.setItem(r, 5, QTableWidgetItem(e.nro_serie or ""))
            self.table.setItem(r, 6, QTableWidgetItem(e.estado or ""))
        if was_sorting:
            self.table.setSortingEnabled(True)
        # Forzar renumeración final para asegurar orden 1..N
        self.update_row_numbers()

    def on_generate(self):
        # Solo usuarios NO administradores deben solicitar permisos extra
        if session.CURRENT_USER != "DI-ADMIN":
            auth = AdminAuthDialog(self, "Para generar reportes de equipos se requieren permisos de administrador.")
            if auth.exec() != QDialog.DialogCode.Accepted:
                QMessageBox.information(self, "Permisos", "Operación cancelada. No se generó el reporte.")
                return

        dlg = GenerateDialog('equipos')
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        vals = dlg.values()
        # Build filtered dataset
        data = list_equipos_by_direccion(self.direccion_filter) if self.direccion_filter else list_equipos()
        def ok(s, sub):
            return (not sub) or (sub.lower() in (s or '').lower())
        filtered = []
        for e in data:
            if not ok(e.codigo_interno, vals.get('codigo')): continue
            if not ok(e.descripcion, vals.get('descripcion')): continue
            if not ok(e.marca, vals.get('marca')): continue
            if not ok(e.modelo, vals.get('modelo')): continue
            if not ok(e.nro_serie, vals.get('serie')): continue
            estado_sel = vals.get('estado') or ""
            if estado_sel and (e.estado or "") != estado_sel:
                continue
            filtered.append(e)
        t = (vals.get('type') or '').lower()
        if t == 'pdf':
            self.generar_pdf_equipos(filtered)
        elif t == 'word':
            self.generar_word_equipos(filtered)
        else:
            self.generar_excel_equipos(filtered)

    def on_table_context_menu(self, pos):
        menu = QMenu(self)
        # Apply dark theme specifically to this menu in case inheritance fails
        menu.setStyleSheet("""
            QMenu { background-color: #1E293B; color: #e0e0e0; border: 1px solid #475569; }
            QMenu::item { padding: 4px 20px; }
            QMenu::item:selected { background-color: #334155; }
        """)
        gen = menu.addMenu("Generar selección")
        # Ensure submenu also uses the style, usually inherits, but let's be safe
        gen.setStyleSheet("""
             QMenu { background-color: #1E293B; color: #e0e0e0; border: 1px solid #475569; }
             QMenu::item { padding: 4px 20px; }
             QMenu::item:selected { background-color: #334155; }
        """)
        
        act_pdf = gen.addAction("PDF")
        act_word = gen.addAction("Word")
        act_excel = gen.addAction("Excel")
        menu.addSeparator()
        act_view = menu.addAction("Ver mantenimientos del equipo")
        action = menu.exec(self.table.viewport().mapToGlobal(pos))
        if action is None:
            return
        # Recoger selección
        selected = []
        data = list_equipos_by_direccion(self.direccion_filter) if self.direccion_filter else list_equipos()
        selected_codes = set()
        for idx in self.table.selectionModel().selectedRows(0):
            it = self.table.item(idx.row(), 1)
            if it:
                selected_codes.add(it.text())
        for e in data:
            if e.codigo_interno in selected_codes:
                selected.append(e)
        if action == act_view:
            if not selected:
                return
            equipo = selected[0]
            # Subir hasta ModulosTab
            parent = self.parent()
            while parent is not None and parent.__class__.__name__ != 'ModulosTab':
                parent = getattr(parent, 'parent', lambda: None)()
            if parent is None:
                return
            # Encontrar índice del tab de Mantenimiento y widget
            mant_widget = None
            mant_index = None
            for i in range(parent.stack.count()):
                w = parent.stack.widget(i)
                if w.__class__.__name__ == 'MantenimientoTab':
                    mant_widget = w
                    mant_index = i
                    break
            if mant_widget is None or mant_index is None:
                return
            # Cambiar a tab de Mantenimiento
            parent.on_nav_clicked(mant_index)
            # Prefijar búsqueda con datos del equipo para filtrar
            term = " ".join(filter(None, [
                equipo.codigo_interno or "",
                equipo.descripcion or "",
                equipo.marca or "",
                equipo.modelo or "",
                equipo.nro_serie or "",
            ])).strip()
            mant_widget.search.setText(term)
            return
        if not selected:
            return

        # Solo usuarios NO administradores deben solicitar permisos extra
        if session.CURRENT_USER != "DI-ADMIN":
            auth = AdminAuthDialog(self, "Para generar reportes de equipos se requieren permisos de administrador.")
            if auth.exec() != QDialog.DialogCode.Accepted:
                QMessageBox.information(self, "Permisos", "Operación cancelada. No se generó el reporte.")
                return

        # Ejecutar según tipo
        if action == act_pdf:
            self.generar_pdf_equipos(selected)
        elif action == act_word:
            self.generar_word_equipos(selected)
        elif action == act_excel:
            self.generar_excel_equipos(selected)

    def generar_pdf_equipos(self, data=None):
        from PyQt6.QtPrintSupport import QPrinter
        from PyQt6.QtGui import QTextDocument
        if data is None:
            data = list_equipos_by_direccion(self.direccion_filter) if self.direccion_filter else list_equipos()
        from datetime import datetime
        fecha_gen = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        html = f"""
        <html>
        <head>
        <meta charset='utf-8' />
        <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #2c3e50; }}
        .header {{ display: flex; align-items: center; justify-content: space-between; margin-bottom: 20px; }}
        .header .info {{ text-align: right; font-size: 12px; color: #7f8c8d; }}
        h1 {{ margin: 0; font-size: 24px; letter-spacing: 0.5px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
        th {{ background-color: #34495e; color: #ecf0f1; font-weight: 600; padding: 10px; font-size: 12px; text-transform: uppercase; }}
        td {{ padding: 8px 10px; font-size: 11px; border-bottom: 1px solid #ecf0f1; }}
        tr:nth-child(even) td {{ background-color: #f8fafc; }}
        .summary {{ margin-top: 20px; font-weight: bold; text-align: right; }}
        </style>
        </head>
        <body>
        <div class='header'>
            <h1>Reporte de Equipos</h1>
            <div class='info'>
                <div>Dirección: {self.header.text().replace('Dirección: ', '') if hasattr(self, 'header') else ''}</div>
                <div>Generado el: {fecha_gen}</div>
            </div>
        </div>
        <table>
        <tr><th>N°</th><th>Código</th><th>Descripción</th><th>Marca</th><th>Modelo</th><th>Serie</th><th>Estado</th></tr>
        """
        for i, e in enumerate(data, 1):
            html += f"<tr><td>{i}</td><td>{e.codigo_interno}</td><td>{e.descripcion or ''}</td><td>{e.marca or ''}</td><td>{e.modelo or ''}</td><td>{e.nro_serie or ''}</td><td>{e.estado or ''}</td></tr>"
        total = len(data)
        html += f"</table><div class='summary'>Total de equipos: {total}</div></body></html>"
        doc = QTextDocument(); doc.setHtml(html)
        printer = QPrinter(); printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        default = f"equipos_direccion_{self.direccion_filter or 'todas'}.pdf"
        fn, _ = QFileDialog.getSaveFileName(self, "Guardar PDF Equipos", default, "PDF (*.pdf)")
        if fn:
            printer.setOutputFileName(fn)
            doc.print(printer)
            dir_name = direccion_nombre(self.direccion_filter)
            add_log("Generar PDF Equipos", f"Archivo: {fn}", dir_name)
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar PDF Equipos", f"Archivo: {fn} - {dir_name if dir_name else 'Todas'}", "Equipos")
            except: pass
            
            QMessageBox.information(self, "PDF Generado", f"PDF guardado en {fn}")

    def generar_word_equipos(self, data=None):
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            QMessageBox.warning(self, "Librería faltante", "Instala python-docx: pip install python-docx")
            return
        if data is None:
            data = list_equipos_by_direccion(self.direccion_filter) if self.direccion_filter else list_equipos()
        from datetime import datetime
        fecha_gen = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = Document(); style = doc.styles['Normal']; style.font.name = 'Segoe UI'; style.font.size = Pt(10.5)
        title = doc.add_heading('Reporte de Equipos', 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.runs[0].font.size = Pt(20); title.runs[0].font.bold = True
        info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info.add_run("Dirección: ").bold = True; info.add_run(self.header.text().replace('Dirección: ', '') if hasattr(self, 'header') else '')
        info.add_run("\nGenerado el: ").bold = True; info.add_run(fecha_gen)
        table = doc.add_table(rows=1, cols=7); table.style = 'Table Grid'; table.autofit = True
        hdr_cells = table.rows[0].cells; headers = ['N°', 'Código', 'Descripción', 'Marca', 'Modelo', 'Serie', 'Estado']
        def _set_cell_shading(cell, fill):
            tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill); tc_pr.append(shd)
        header_color = RGBColor(0xFF, 0xFF, 0xFF)
        for i, h in enumerate(headers):
            paragraph = hdr_cells[i].paragraphs[0]; paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(h); run.bold = True; run.font.color.rgb = header_color; _set_cell_shading(hdr_cells[i], '34495E'); hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i, e in enumerate(data, 1):
            row_cells = table.add_row().cells
            values = [str(i), str(e.codigo_interno or ''), str(e.descripcion or ''), str(e.marca or ''), str(e.modelo or ''), str(e.nro_serie or ''), str(e.estado or '')]
            for idx, value in enumerate(values):
                cell = row_cells[idx]; cell.text = value; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if i % 2 == 0: _set_cell_shading(cell, 'F8FAFC')
        total_row = table.add_row().cells; total_row[0].merge(total_row[6])
        total_paragraph = total_row[0].paragraphs[0]; total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_run = total_paragraph.add_run(f"Total de equipos: {len(data)}"); total_run.bold = True
        for section in doc.sections:
            section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5); section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
        fn, _ = QFileDialog.getSaveFileName(self, "Guardar Word Equipos", f"equipos_direccion_{self.direccion_filter or 'todas'}.docx", "Word (*.docx)")
        if fn:
            doc.save(fn)
            dir_name = direccion_nombre(self.direccion_filter)
            add_log("Generar Word Equipos", f"Archivo: {fn}", dir_name)
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar Word Equipos", f"Archivo: {fn} - {dir_name if dir_name else 'Todas'}", "Equipos")
            except: pass

            QMessageBox.information(self, "Word Generado", f"Documento guardado en {fn}")

    def generar_excel_equipos(self, data=None):
        if data is None:
            data = list_equipos_by_direccion(self.direccion_filter) if self.direccion_filter else list_equipos()
        from datetime import datetime
        fecha_gen = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        temp_table = QTableWidget(0, 6)
        temp_table.setHorizontalHeaderLabels(["Código", "Descripción", "Marca", "Modelo", "Serie", "Estado"])
        for e in data:
            r = temp_table.rowCount(); temp_table.insertRow(r)
            temp_table.setItem(r, 0, QTableWidgetItem(e.codigo_interno or ''))
            temp_table.setItem(r, 1, QTableWidgetItem(e.descripcion or ''))
            temp_table.setItem(r, 2, QTableWidgetItem(e.marca or ''))
            temp_table.setItem(r, 3, QTableWidgetItem(e.modelo or ''))
            temp_table.setItem(r, 4, QTableWidgetItem(e.nro_serie or ''))
            temp_table.setItem(r, 5, QTableWidgetItem(e.estado or ''))
        fn, _ = QFileDialog.getSaveFileName(self, "Guardar Excel Equipos", f"equipos_direccion_{self.direccion_filter or 'todas'}.xlsx", "Excel (*.xlsx)")
        if fn:
            export_table_to_excel(temp_table, fn)
            dir_name = direccion_nombre(self.direccion_filter)
            add_log("Generar Excel Equipos", f"Archivo: {fn}", dir_name)
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar Excel Equipos", f"Archivo: {fn} - {dir_name if dir_name else 'Todas'}", "Equipos")
            except: pass

            QMessageBox.information(self, "Excel Generado", f"Archivo guardado en {fn}")

    def show_details(self, *_):
        id_ = self.current_id()
        if not id_:
            return
        equipo = get_equipo(id_)
        if not equipo:
            return
        dir_name = direccion_nombre(equipo.direccion_id)
        details = [
            ("Código", equipo.codigo_interno or ""),
            ("Descripción", equipo.descripcion or ""),
            ("Marca", equipo.marca or ""),
            ("Modelo", equipo.modelo or ""),
            ("Serie", equipo.nro_serie or ""),
            ("Estado", equipo.estado or ""),
            ("Dirección", dir_name or ""),
        ]
        dialog = RecordDetailDialog("Detalle de Equipo", details)
        dialog.exec()

    def current_id(self):
        r = self.table.currentRow()
        if r < 0:
            return None
        # Buscar el ID por código interno (ahora en columna 1)
        codigo_sel = self.table.item(r, 1).text()
        data = list_equipos_by_direccion(self.direccion_filter) if self.direccion_filter else list_equipos()
        for e in data:
            if e.codigo_interno == codigo_sel:
                return e.id
        return None

    def _update_header(self):
        if not self.direccion_filter:
            self.header.setText("Dirección: (todas)")
            return
        name = ""
        try:
            for d in list_direcciones():
                if d.id == self.direccion_filter:
                    name = d.nombre
                    break
        except Exception:
            name = ""
        self.header.setText(f"{name}")

    def on_add_modal(self):
        dlg = EquipoDialog(direccion_filter=self.direccion_filter)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        vals = dlg.values()
        if not vals["codigo_interno"] or not vals["descripcion"]:
            QMessageBox.information(self, "Datos requeridos", "Complete Código y Descripción.")
            return
        if self.direccion_filter:
            vals["direccion_id"] = self.direccion_filter
        try:
            add_equipo({
                **vals,
                "fecha_alta": date.today(),
            })
        except IntegrityError:
            QMessageBox.warning(self, "Código duplicado", "Ya existe un Equipo con ese código interno.")
            return
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo agregar el equipo.\n{e}")
            return
        self.refresh()
        # seleccionar el creado
        try:
            for r in range(self.table.rowCount()):
                it = self.table.item(r, 1)
                if it and it.text() == vals["codigo_interno"]:
                    self.table.selectRow(r)
                    self.on_select(r, 0)
                    break
        except Exception:
            pass
        dir_name = direccion_nombre(vals.get("direccion_id") or self.direccion_filter)
        add_log("Agregar Equipo", f"Código: {vals['codigo_interno']}, Descripción: {vals['descripcion']}", dir_name)
        
        # Bitacora
        try:
            u_obj = get_user(session.CURRENT_USER)
            if u_obj:
                add_bitacora_log(u_obj.id, "Agregar Equipo", f"Equipo: {vals['codigo_interno']} en {dir_name}", "Equipos")
        except: pass

    def on_edit_modal(self):
        id_ = self.current_id()
        if not id_:
            return
        # reunir datos actuales desde la tabla
        r = self.table.currentRow()
        cur = {
            "codigo_interno": self.table.item(r, 1).text() if self.table.item(r, 1) else "",
            "descripcion": self.table.item(r, 2).text() if self.table.item(r, 2) else "",
            "marca": self.table.item(r, 3).text() if self.table.item(r, 3) else "",
            "modelo": self.table.item(r, 4).text() if self.table.item(r, 4) else "",
            "nro_serie": self.table.item(r, 5).text() if self.table.item(r, 5) else "",
            "estado": self.table.item(r, 6).text() if self.table.item(r, 6) else "optimo",
            "direccion_id": self.direccion_filter or (self.dir.currentData() if hasattr(self, 'dir') else None),
        }
        dlg = EquipoDialog(direccion_filter=self.direccion_filter, data=cur)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        vals = dlg.values()
        if self.direccion_filter:
            vals["direccion_id"] = self.direccion_filter
        try:
            old = get_equipo(id_)
            dir_name = direccion_nombre(vals.get("direccion_id") or old.direccion_id or self.direccion_filter)
            current_codigo = vals.get("codigo_interno", old.codigo_interno)
            changes = []
            for k, v in vals.items():
                old_v = getattr(old, k)
                if old_v != v:
                    field_name = {
                        'codigo_interno': 'Código',
                        'descripcion': 'Desc',
                        'marca': 'Marca',
                        'modelo': 'Modelo',
                        'nro_serie': 'Serie',
                        'estado': 'Estado',
                        'direccion_id': 'Dirección ID'
                    }.get(k, k)
                    changes.append(f"{field_name}: '{old_v}'->'{v}'")

            if changes:
                change_str = ", ".join(changes)
                add_log("Editar Equipo", f"Equipo {current_codigo}: {change_str}", dir_name)
                
                # Bitacora DB
                try:
                    u_obj = get_user(session.CURRENT_USER)
                    if u_obj:
                        add_bitacora_log(u_obj.id, "Editar Equipo", f"Equipo {current_codigo}: {change_str}", "Equipos")
                except: pass

            update_equipo(id_, vals)
            
        except IntegrityError:
            QMessageBox.warning(self, "Código duplicado", "Ya existe un Equipo con ese código interno.")
            return
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo actualizar el equipo.\n{e}")
            return
        self.refresh()
        grandparent = self.parent().parent()
        if hasattr(grandparent, 'refresh'):
            grandparent.refresh()

    def on_delete(self):
        id_ = self.current_id()
        if not id_:
            QMessageBox.information(self, "Selección requerida", "Selecciona un equipo para eliminar.")
            return
        codigo = self.table.item(self.table.currentRow(), 1).text()
        # Solo usuarios NO administradores deben solicitar permisos extra
        if session.CURRENT_USER != "DI-ADMIN":
            auth = AdminAuthDialog(self, "Para eliminar equipos se requieren permisos de administrador.")
            if auth.exec() != QDialog.DialogCode.Accepted:
                QMessageBox.information(self, "Permisos", "Operación cancelada. No se eliminó el equipo.")
                return

        if QMessageBox.question(self, "Confirmar", "¿Eliminar equipo seleccionado?") == QMessageBox.StandardButton.Yes:
            equipo = get_equipo(id_)
            dir_name = direccion_nombre(getattr(equipo, 'direccion_id', None) if equipo else self.direccion_filter)
            delete_equipo(id_)
            add_log("Eliminar Equipo", f"Código: {codigo}", dir_name)
            
            # Bitacora
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Eliminar Equipo", f"Eliminó el equipo: {codigo} de {dir_name}", "Equipos")
            except: pass
            
        self.refresh()

    def on_select(self, r, c):
        pass


from datetime import date, datetime
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton,
    QTableWidget, QTableWidgetItem, QHeaderView, QComboBox, QDateEdit,
    QFileDialog, QMessageBox, QDialog
)
from PyQt6.QtCore import Qt, QDate, QTimer

from ...utils import load_icon, export_table_to_excel, NumericTableWidgetItem
from ...logger import add_log
from ... import session
from ..helpers import direccion_nombre
from ..dialogs import GenerateDialog, MantenimientoDialog, RecordDetailDialog, AdminAuthDialog
from ...data.repositories import (
    list_direcciones, list_equipos, list_mantenimientos, 
    list_mantenimientos_by_direccion, add_mantenimiento,
    update_mantenimiento, delete_mantenimiento, get_mantenimiento,
    get_equipo, update_equipo,
    add_bitacora_log, get_user
)
from sqlalchemy.exc import IntegrityError

class MantenimientoTab(QWidget):
    def __init__(self, departamento_id: int | None = None, direccion_id: int | None = None):
        super().__init__()
        self.departamento_filter = None
        self.direccion_filter = direccion_id
        self.table = QTableWidget(0, 7)
        self.table.setHorizontalHeaderLabels([
            "N°", "Equipo (Cód)","Desc. Equipo","Fecha","Observación","Estado Eq.","Dirección"
        ])
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.table.setAlternatingRowColors(True)
        self.table.setShowGrid(False)  # Desactivar grid nativo para estilo CSS premium
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents) # Columna N°
        self.table.setSortingEnabled(True)

        self.search = QLineEdit(); self.search.setPlaceholderText("Buscar por código/descripción/obs…")
        self.de_from = QDateEdit(); self.de_from.setDate(QDate.currentDate().addMonths(-1))
        self.de_to = QDateEdit(); self.de_to.setDate(QDate.currentDate())
        self.cb_estado = QComboBox(); self.cb_estado.addItems(["", "optimo","defectuoso","inoperativo"])

        self.btn_add = QPushButton("Nuevo")
        self.btn_upd = QPushButton("Editar")
        self.btn_del = QPushButton("Eliminar")
        self.btn_add.clicked.connect(self.on_add)
        self.btn_upd.clicked.connect(self.on_edit)
        self.btn_del.clicked.connect(self.on_delete)

        # Layout Principal
        layout = QHBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 24)
        layout.setSpacing(20)

        # --- Panel Izquierdo: Acciones ---
        actions_card = QWidget()
        actions_card.setProperty("class", "card")
        actions_card.setFixedWidth(200)
        
        actions_l = QVBoxLayout(actions_card)
        actions_l.setContentsMargins(16, 20, 16, 20)
        actions_l.setSpacing(12)
        
        lbl_actions = QLabel("Acciones")
        lbl_actions.setStyleSheet("font-weight: 700; color: #94A3B8; text-transform: uppercase; font-size: 12px; margin-bottom: 8px;")
        actions_l.addWidget(lbl_actions)

        # Botones
        for btn, label, icon in [
            (self.btn_add, "Nuevo Reg.", "plus.svg"),
            (self.btn_upd, "Editar Selecc.", "edit.svg"),
            (self.btn_del, "Eliminar Selecc.", "trash.svg")
        ]:
            btn.setText(label)
            btn.setIcon(load_icon(icon))
            btn.setProperty("class", "panel-accent")
            btn.setMinimumHeight(42)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            if btn == self.btn_del:
                 btn.setStyleSheet("QPushButton { color: #EF4444; border: 1px solid rgba(239, 68, 68, 0.3); } QPushButton:hover { background: rgba(239, 68, 68, 0.1); border: 1px solid rgba(239, 68, 68, 0.8); }")
            actions_l.addWidget(btn)

        gen_btn = QPushButton("Generar Reporte")
        gen_btn.setProperty("class", "panel-accent")
        gen_btn.setMinimumHeight(42)
        gen_btn.setIcon(load_icon("records.svg"))
        gen_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        gen_btn.clicked.connect(self.on_generate)
        
        actions_l.addStretch(1)
        actions_l.addWidget(gen_btn)
        
        layout.addWidget(actions_card)

        # --- Panel Derecho: Contenido ---
        content_layout = QVBoxLayout()
        content_layout.setSpacing(16)
        
        # Header Row
        header_row = QHBoxLayout()
        
        self.header = QLabel("")
        self.header.setStyleSheet("font-size: 24px; font-weight: 700; color: #F8FAFC;")
        
        self.search.setPlaceholderText("Buscar mantenimientos...")
        self.search.setFixedWidth(300)
        self.search.setStyleSheet("""
            QLineEdit {
                border-radius: 18px;
                padding: 8px 16px;
                background: rgba(30, 41, 59, 0.6);
                border: 1px solid rgba(148, 163, 184, 0.2);
                color: #F1F5F9;
            }
            QLineEdit:focus {
                border: 1px solid #38BDF8;
                background: rgba(30, 41, 59, 0.9);
            }
        """)

        header_row.addWidget(self.header)
        header_row.addStretch(1)
        header_row.addWidget(self.search)
        
        content_layout.addLayout(header_row)

        # Tabla Card
        table_card = QWidget()
        table_card.setProperty("class", "card")
        table_l = QVBoxLayout(table_card)
        table_l.setContentsMargins(0, 0, 0, 0)
        
        self.table.setShowGrid(False)
        self.table.verticalHeader().setVisible(False)
        self.table.setFrameShape(QTableWidget.Shape.NoFrame)
        
        table_l.addWidget(self.table)
        content_layout.addWidget(table_card, 1)

        layout.addLayout(content_layout, 1)

        # Init logic
        self._update_header()
        self.refresh()
        self.search.textChanged.connect(self.refresh)
        self.table.itemDoubleClicked.connect(self.on_edit)
        # Auto-renumerar al ordenar
        self.table.horizontalHeader().sectionClicked.connect(self.schedule_renumber)
        self.table.horizontalHeader().sortIndicatorChanged.connect(self.schedule_renumber)

    def schedule_renumber(self):
        QTimer.singleShot(50, self.update_row_numbers)

    def update_row_numbers(self):
        if self.table.model(): self.table.model().blockSignals(True)
        for r in range(self.table.rowCount()):
            item = self.table.item(r, 0)
            if not item:
                item = NumericTableWidgetItem(str(r + 1))
                self.table.setItem(r, 0, item)
            else:
                item.setText(str(r + 1))
            item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        if self.table.model(): self.table.model().blockSignals(False)

    def _update_header(self):
        if not self.direccion_filter:
            self.header.setText("Mantenimientos: Todas las Direcciones")
            return
        name = direccion_nombre(self.direccion_filter)
        self.header.setText(f"Dirección: {name}")

    def refresh(self):
        if self.direccion_filter:
            data = list_mantenimientos_by_direccion(self.direccion_filter)
        else:
            data = list_mantenimientos()
        was_sorting = self.table.isSortingEnabled()
        if was_sorting:
            self.table.setSortingEnabled(False)
        self.table.setRowCount(0)
        term = self.search.text().strip().lower()
        for m in data:
            eq = m.equipo
            eq_code = eq.codigo_interno if eq else ""
            eq_desc = eq.descripcion if eq else ""
            obs = m.descripcion or ""
            date_str = str(m.fecha)
            est = m.estado_equipo or ""
            d_id = eq.direccion_id if eq else None
            dir_name = direccion_nombre(d_id)
            blob = f"{eq_code} {eq_desc} {obs} {date_str} {est} {dir_name}".lower()
            if term and term not in blob:
                continue
            r = self.table.rowCount()
            r = self.table.rowCount()
            self.table.insertRow(r)
            # N° (y guardamos ID aqui)
            item_n = NumericTableWidgetItem(str(r + 1))
            item_n.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.table.setItem(r, 0, item_n)
            self.table.item(r, 0).setData(Qt.ItemDataRole.UserRole, m.id)

            self.table.setItem(r, 1, QTableWidgetItem(eq_code))
            self.table.setItem(r, 2, QTableWidgetItem(eq_desc))
            self.table.setItem(r, 3, QTableWidgetItem(date_str))
            self.table.setItem(r, 4, QTableWidgetItem(obs))
            self.table.setItem(r, 5, QTableWidgetItem(est))
            self.table.setItem(r, 6, QTableWidgetItem(dir_name))
        if was_sorting:
            self.table.setSortingEnabled(True)
        self.update_row_numbers()

    def current_id(self):
        r = self.table.currentRow()
        if r < 0:
            return None
        return self.table.item(r, 0).data(Qt.ItemDataRole.UserRole)

    def on_add(self):
        dlg = MantenimientoDialog(direccion_filter=self.direccion_filter)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        vals = dlg.values()
        if not vals["equipo_id"]:
            QMessageBox.information(self, "Requerido", "Seleccione un equipo.")
            return
        try:
            old_eq = get_equipo(vals["equipo_id"])
            if old_eq:
                m_estado = vals.get("estado_equipo")
                # update equipo state
                if m_estado and old_eq.estado != m_estado:
                    update_equipo(old_eq.id, {"estado": m_estado})
                    add_log("Actualización automática de estado", f"Equipo {old_eq.codigo_interno} pasa a {m_estado}", direccion_nombre(old_eq.direccion_id))
            add_mantenimiento(vals)
            dir_name = direccion_nombre(old_eq.direccion_id) if old_eq else ""
            add_log("Agregar Mantenimiento", f"Equipo: {old_eq.codigo_interno if old_eq else '?'}", dir_name)
            
            # Bitacora
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Agregar Mantenimiento", f"Equipo: {old_eq.codigo_interno if old_eq else 'unknown'} en {dir_name}", "Mantenimientos")
            except: pass
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Fallo al agregar: {e}")
        self.refresh()

    def on_edit(self):
        id_ = self.current_id()
        if not id_:
            return
        mant = get_mantenimiento(id_)
        if not mant:
            return
        data = {
            "equipo_id": mant.equipo_id,
            "fecha": mant.fecha,
            "descripcion": mant.descripcion,
            "estado_equipo": mant.estado_equipo,
        }
        dlg = MantenimientoDialog(direccion_filter=self.direccion_filter, data=data)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        vals = dlg.values()
        try:
            eq = get_equipo(vals['equipo_id'])
            # Detect changes for log
            change_details = []
            if mant.equipo_id != vals['equipo_id']: change_details.append(f"Equipo: {mant.equipo_id}->{vals['equipo_id']}")
            if str(mant.fecha) != str(vals['fecha']): change_details.append(f"Fecha: {mant.fecha}->{vals['fecha']}")
            if mant.descripcion != vals['descripcion']: change_details.append(f"Desc: '{mant.descripcion}'->'{vals['descripcion']}'")
            if mant.estado_equipo != vals['estado_equipo']: change_details.append(f"Estado: {mant.estado_equipo}->{vals['estado_equipo']}")
            
            dir_name = direccion_nombre(eq.direccion_id) if eq else ""
            if change_details:
                details_text = ", ".join(change_details)
                add_log("Editar Mantenimiento", f"Equipo {eq.codigo_interno if eq else '?'}: {details_text}", dir_name)
                
                # Bitacora DB con detalles
                try:
                    u_obj = get_user(session.CURRENT_USER)
                    if u_obj:
                        add_bitacora_log(
                            u_obj.id, 
                            "Editar Mantenimiento", 
                            f"Equipo {eq.codigo_interno if eq else '?'}: {details_text}", 
                            "Mantenimientos"
                        )
                except Exception: pass

            update_mantenimiento(id_, vals)
            # update equipo state if changed
            if eq and vals.get("estado_equipo") and eq.estado != vals["estado_equipo"]:
                update_equipo(eq.id, {"estado": vals["estado_equipo"]})
                add_log("Actualización automática de estado", f"Equipo {eq.codigo_interno} pasa a {vals['estado_equipo']}", dir_name)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Fallo al actualizar: {e}")
        self.refresh()

    def on_delete(self):
        id_ = self.current_id()
        if not id_:
            QMessageBox.information(self, "Selección", "Elige un mantenimiento.")
            return
        # Solo usuarios NO administradores deben solicitar permisos extra
        if session.CURRENT_USER != "DI-ADMIN":
            auth = AdminAuthDialog(self, "Para eliminar registros de mantenimiento se requieren permisos de administrador.")
            if auth.exec() != QDialog.DialogCode.Accepted:
                QMessageBox.information(self, "Permisos", "Operación cancelada. No se eliminó el registro.")
                return

        if QMessageBox.question(self, "Confirmar", "¿Eliminar registro de mantenimiento?") == QMessageBox.StandardButton.Yes:
            mant = get_mantenimiento(id_)
            eq_code = mant.equipo.codigo_interno if mant and mant.equipo else "?"
            dir_name = direccion_nombre(mant.equipo.direccion_id) if mant and mant.equipo else ""
            delete_mantenimiento(id_)
            add_log("Eliminar Mantenimiento", f"Del equipo {eq_code}", dir_name)
            
            # Bitacora
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Eliminar Mantenimiento", f"Del equipo {eq_code} en {dir_name}", "Mantenimientos")
            except: pass
            
            self.refresh()

    def on_generate(self):
        # Solo usuarios NO administradores deben solicitar permisos extra
        if session.CURRENT_USER != "DI-ADMIN":
            auth = AdminAuthDialog(self, "Para generar reportes de mantenimientos se requieren permisos de administrador.")
            if auth.exec() != QDialog.DialogCode.Accepted:
                QMessageBox.information(self, "Permisos", "Operación cancelada. No se generó el reporte.")
                return

        dlg = GenerateDialog('mantenimientos')
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        vals = dlg.values()
        data = list_mantenimientos_by_direccion(self.direccion_filter) if self.direccion_filter else list_mantenimientos()
        # Filter
        f_from = vals.get('from')
        f_to = vals.get('to')
        f_est = vals.get('estado')
        f_eq = vals.get('equipo')
        f_obs = vals.get('obs')

        filtered = []
        for m in data:
            if f_from and str(m.fecha) < f_from: continue
            if f_to and str(m.fecha) > f_to: continue
            if f_est and m.estado_equipo != f_est: continue
            eq = m.equipo
            if f_eq:
                blob = f"{eq.codigo_interno or ''} {eq.descripcion or ''} {eq.marca or ''} {eq.modelo or ''} {eq.nro_serie or ''}".lower()
                if f_eq.lower() not in blob: continue
            if f_obs and f_obs.lower() not in (m.descripcion or '').lower(): continue
            filtered.append(m)

        t = vals.get('type', '').lower()
        if t == 'pdf': self.generar_pdf_mantenimientos(filtered)
        elif t == 'word': self.generar_word_mantenimientos(filtered)
        else: self.generar_excel_mantenimientos(filtered)

    def generar_pdf_mantenimientos(self, data):
        from PyQt6.QtPrintSupport import QPrinter
        from PyQt6.QtGui import QTextDocument
        fecha_gen = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        html = f"""
        <html>
        <head>
        <meta charset='utf-8' />
        <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #2c3e50; }}
        .header {{ display: flex; align-items: center; justify-content: space-between; margin-bottom: 20px; }}
        .header .info {{ text-align: right; font-size: 12px; color: #7f8c8d; }}
        h1 {{ margin: 0; font-size: 24px; letter-spacing: 0.5px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
        th {{ background-color: #34495e; color: #ecf0f1; font-weight: 600; padding: 10px; font-size: 12px; text-transform: uppercase; }}
        td {{ padding: 8px 10px; font-size: 11px; border-bottom: 1px solid #ecf0f1; }}
        tr:nth-child(even) td {{ background-color: #f8fafc; }}
        .summary {{ margin-top: 20px; font-weight: bold; text-align: right; }}
        </style>
        </head>
        <body>
        <div class='header'>
            <h1>Reporte de Mantenimientos</h1>
            <div class='info'>
                <div>{self.header.text() if hasattr(self, 'header') else ''}</div>
                <div>Generado el: {fecha_gen}</div>
            </div>
        </div>
        <table>
        <tr><th>N°</th><th>Equipo</th><th>Fecha</th><th>Observación</th><th>Estado Eq.</th><th>Dirección</th></tr>
        """
        for i, m in enumerate(data, 1):
            eq = m.equipo
            eq_str = f"{eq.codigo_interno or ''} ({eq.descripcion or ''})" if eq else "-"
            dir_name = direccion_nombre(eq.direccion_id) if eq else ""
            html += f"<tr><td>{i}</td><td>{eq_str}</td><td>{m.fecha}</td><td>{m.descripcion or ''}</td><td>{m.estado_equipo or ''}</td><td>{dir_name}</td></tr>"
        html += f"</table><div class='summary'>Registros: {len(data)}</div></body></html>"
        
        doc = QTextDocument(); doc.setHtml(html)
        printer = QPrinter(); printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        default = "reporte_mantenimientos.pdf"
        fn, _ = QFileDialog.getSaveFileName(self, "PDF", default, "PDF (*.pdf)")
        if fn:
            printer.setOutputFileName(fn)
            doc.print(printer)
            dir_name = direccion_nombre(self.direccion_filter)
            add_log("Generar PDF Mantenimientos", f"Archivo: {fn}", dir_name)
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar PDF Mantenimientos", f"Archivo: {fn} - {dir_name if dir_name else 'Todas'}", "Mantenimientos")
            except: pass

            QMessageBox.information(self, "OK", f"Guardado en {fn}")

    def generar_word_mantenimientos(self, data):
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except:
            QMessageBox.warning(self, "Error", "Instalar python-docx")
            return
        fecha_gen = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = Document(); style = doc.styles['Normal']; style.font.name = 'Segoe UI'; style.font.size = Pt(10.5)
        title = doc.add_heading('Reporte de Mantenimientos', 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info.add_run(self.header.text() + "\n").bold = True
        info.add_run("Generado: " + fecha_gen)
        
        table = doc.add_table(rows=1, cols=6); table.style = 'Table Grid'; table.autofit = True
        headers = ['N°','Equipo','Fecha','Obs', 'Est. Eq.', 'Dirección']
        def _shade(c, color):
            tc = c._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color); tc.append(shd)
        for i, h in enumerate(headers):
            p = table.rows[0].cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h); r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
            _shade(table.rows[0].cells[i], "34495E")
            
        for i, m in enumerate(data, 1):
            eq = m.equipo
            eq_str = f"{eq.codigo_interno}" if eq else "-"
            dir_name = direccion_nombre(eq.direccion_id) if eq else ""
            cells = table.add_row().cells
            vals = [str(i), eq_str, str(m.fecha), str(m.descripcion or ''), str(m.estado_equipo or ''), dir_name]
            for j, v in enumerate(vals): cells[j].text = v
            if i % 2 == 0:
                for c in cells: _shade(c, "F8FAFC")
        
        default = "reporte_mantenimientos.docx"
        fn, _ = QFileDialog.getSaveFileName(self, "Word", default, "Word (*.docx)")
        if fn:
            doc.save(fn)
            dir_name = direccion_nombre(self.direccion_filter)
            add_log("Generar Word Mantenimientos", f"Archivo: {fn}", dir_name)
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar Word Mantenimientos", f"Archivo: {fn} - {dir_name if dir_name else 'Todas'}", "Mantenimientos")
            except: pass
            
            QMessageBox.information(self, "OK", f"Guardado en {fn}")

    def generar_excel_mantenimientos(self, data):
        rows = []
        for m in data:
            eq = m.equipo
            dir_name = direccion_nombre(eq.direccion_id) if eq else ""
            rows.append([
                eq.codigo_interno if eq else "", eq.descripcion if eq else "", 
                str(m.fecha), m.descripcion or "", m.estado_equipo or "", dir_name
            ])
        tmp = QTableWidget(len(rows), 6)
        tmp.setHorizontalHeaderLabels(["Equipo","Desc. Equipo","Fecha","Obs","Estado","Dirección"])
        for i, r in enumerate(rows):
            for j, v in enumerate(r):
                tmp.setItem(i, j, QTableWidgetItem(v))
        default = "reporte_mantenimientos.xlsx"
        fn, _ = QFileDialog.getSaveFileName(self, "Excel", default, "Excel (*.xlsx)")
        if fn:
            export_table_to_excel(tmp, fn)
            dir_name = direccion_nombre(self.direccion_filter)
            add_log("Generar Excel Mantenimientos", f"Archivo: {fn}", dir_name)
            
            # Bitacora DB
            try:
                u_obj = get_user(session.CURRENT_USER)
                if u_obj:
                    add_bitacora_log(u_obj.id, "Generar Excel Mantenimientos", f"Archivo: {fn} - {dir_name if dir_name else 'Todas'}", "Mantenimientos")
            except: pass

            QMessageBox.information(self, "OK", f"Guardado en {fn}")
